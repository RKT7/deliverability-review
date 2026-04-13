import docx
from io import BytesIO


def read_docx(file):
    """Extract all text from a .docx file.
    
    Args:
        file: Streamlit UploadedFile object
        
    Returns:
        str: Full text content of the document
    """
    doc = docx.Document(BytesIO(file.read()))
    
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # Also extract from tables (RFPs often have tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    
    return '\n'.join(paragraphs)
